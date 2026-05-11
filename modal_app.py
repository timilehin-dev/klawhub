import modal
import subprocess
import tempfile
import os
import json
import resource
import base64
import glob
import hmac
import hashlib
import time
import requests
from fastapi import Request, HTTPException
from bs4 import BeautifulSoup

app = modal.App("klawhub-sandbox")

# ── Secrets ──
# Create this secret first:  modal secret create klawhub-webhook-secret
# Then set MODAL_WEBHOOK_SECRET to the same value in your Vercel env.
webhook_secret = modal.Secret.from_name("klawhub-webhook-secret")

image = (
    modal.Image.debian_slim(python_version="3.11")
    .apt_install("nodejs")  # Enable JavaScript execution
    .pip_install(
        "fastapi[standard]",
        "requests",
        "beautifulsoup4",
        "pandas",
        "numpy",
        "matplotlib",
        "seaborn",
        "python-docx",
        "reportlab",
        "pdfplumber",
        "pypdf",
        "fastembed",
    )
    .run_commands("python -c 'from fastembed import TextEmbedding; TextEmbedding()'")
)


# ─────────────────────────────────────────────
# Auth Middleware
# ─────────────────────────────────────────────

def verify_request(request: Request) -> bool:
    """Verify requests using a shared webhook secret."""
    provided = request.headers.get("X-Webhook-Secret", "")
    expected = modal.container_app.secrets["klawhub-webhook-secret"]

    if not expected:
        # Secret MUST be configured in production
        return False

    if not provided:
        return False

    return hmac.compare_digest(provided, expected)


# ─────────────────────────────────────────────
# Code Execution
# ─────────────────────────────────────────────

@app.function(image=image, timeout=300)
def execute_code(code: str, language: str = "python", dependencies: str = ""):
    if language not in ["python", "javascript"]:
        return {"success": False, "error": f"Unsupported language: {language}"}

    if dependencies:
        # CRITICAL: Arbitrary pip/npm install is disabled for security.
        # Use the pre-configured 'image' to add new persistent dependencies.
        return {
            "success": False,
            "error": "Dynamic dependency installation is disabled for security. Please request image updates."
        }

    try:
        ext = ".py" if language == "python" else ".js"
        cmd = ["python3"] if language == "python" else ["node"]

        with tempfile.NamedTemporaryFile(mode="w", suffix=ext, delete=False) as f:
            f.write(code)
            filepath = f.name

        def set_resource_limits():
            # Set virtual memory limit to 512 MB (512 * 1024 * 1024 bytes).
            # RLIMIT_AS limits the total address space (virtual memory).
            # Both soft and hard limits are set to the same value.
            # These limits are for the child process executing user code.
            # Note: These values are examples and should be tuned based on expected usage.
            try:
                resource.setrlimit(resource.RLIMIT_AS, (512 * 1024 * 1024, 512 * 1024 * 1024))
            except ValueError:
                # This can happen if the limit is too high for the system or already exceeded.
                # In a production environment, this should be logged.
                pass

            # Set CPU time limit to 30 seconds.
            # RLIMIT_CPU limits the amount of CPU time the process can consume.
            # Both soft and hard limits are set to the same value.
            # Note: This value aligns with the subprocess.run timeout.
            try:
                resource.setrlimit(resource.RLIMIT_CPU, (30, 30))
            except ValueError:
                # In a production environment, this should be logged.
                pass

        try:
            result = subprocess.run(
                cmd + [filepath],
                capture_output=True,
                timeout=30, # This timeout is for the parent process waiting for the child.
                preexec_fn=set_resource_limits # Apply resource limits to the child process before exec.
            )
            # Decode stdout/stderr explicitly to avoid UnicodeDecodeError
            stdout_decoded = result.stdout.decode('utf-8', errors='replace')
            stderr_decoded = result.stderr.decode('utf-8', errors='replace')
        finally:
            try:
                os.unlink(filepath)
            except OSError:
                pass

        return {
            "success": result.returncode == 0,
            "stdout": result.stdout[:10000],
            "stderr": result.stderr[:5000],
            "error": None if result.returncode == 0 else f"Exit code {result.returncode}",
        }

    except subprocess.TimeoutExpired:
        return {"success": False, "error": "Code execution timed out (30s)"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ─────────────────────────────────────────────
# Web Page Reader
# ─────────────────────────────────────────────

@app.function(image=image, timeout=30)
def read_web_page(url: str):
    try:
        headers = {"User-Agent": "Klawhub/1.0 (Research Bot)"}
        resp = requests.get(url, headers=headers, timeout=15, verify=True)
        resp.raise_for_status()

        soup = BeautifulSoup(resp.text, "html.parser")

        # Remove scripts, styles, nav, footer, header, aside
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        content = "\n".join(lines)

        return {
            "success": True,
            "content": content[:8000],
            "title": soup.title.string if soup.title else "",
        }

    except Exception as e:
        return {"success": False, "error": str(e)}


# ─────────────────────────────────────────────
# Document Generation
# ─────────────────────────────────────────────

@app.function(image=image, timeout=60)
def generate_document(data: dict):
    format_type = data.get("format", "pdf")
    title = data.get("title", "Document")
    sections = data.get("sections", [])

    try:
        if format_type == "docx":
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            style = doc.styles["Normal"]
            font = style.font
            font.name = "Calibri"
            font.size = Pt(11)

            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for section in sections:
                if section.get("heading"):
                    doc.add_heading(section["heading"], level=1)
                if section.get("body"):
                    doc.add_paragraph(section["body"])

            filepath = f"/tmp/{_safe_filename(title)}_{int(time.time())}.docx"
            doc.save(filepath)

        elif format_type == "pdf":
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

            filepath = f"/tmp/{_safe_filename(title)}_{int(time.time())}.pdf"
            doc = SimpleDocTemplate(
                filepath, pagesize=letter,
                leftMargin=0.75*inch, rightMargin=0.75*inch,
            )

            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "CustomTitle", parent=styles["Title"],
                fontSize=24, spaceAfter=20, alignment=1,
            )
            heading_style = ParagraphStyle(
                "CustomHeading", parent=styles["Heading2"],
                fontSize=14, spaceBefore=16, spaceAfter=8,
            )
            body_style = ParagraphStyle(
                "CustomBody", parent=styles["Normal"],
                fontSize=11, leading=16, spaceAfter=8,
            )

            story = [Paragraph(title, title_style), Spacer(1, 20)]

            for section in sections:
                if section.get("heading"):
                    story.append(Paragraph(section["heading"], heading_style))
                if section.get("body"):
                    story.append(Paragraph(section["body"], body_style))

            doc.build(story)
        else:
            return {"success": False, "error": f"Unsupported format: {format_type}"}

        with open(filepath, "rb") as f:
            file_b64 = base64.b64encode(f.read()).decode()

        filename = os.path.basename(filepath)

        return {
            "success": True,
            "output_file": file_b64,
            "filename": filename,
        }

    except Exception as e:
        return {"success": False, "error": str(e)}


# ─────────────────────────────────────────────
# Data Analytics
# ─────────────────────────────────────────────

def _cleanup_stale_charts(prefix: str = ""):
    """Remove old chart PNGs to avoid stale data from previous runs."""
    pattern = f"/tmp/{prefix}*.png" if prefix else "/tmp/*.png"
    try:
        for f in glob.glob(pattern):
            # Only delete files older than 60 seconds (from previous runs)
            if os.path.getmtime(f) < time.time() - 60:
                os.unlink(f)
    except OSError:
        pass


@app.function(image=image, timeout=120, memory=512, secrets=[webhook_secret])
def run_analytics(data: dict):
    code = data.get("code", "")

    if not code:
        return {"success": False, "error": "No code provided"}

    # Clean stale charts from previous runs
    _cleanup_stale_charts()

    preamble = """
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import io, os, sys, time

plt.rcParams['figure.dpi'] = 150
plt.rcParams['savefig.bbox'] = 'tight'
plt.rcParams['font.size'] = 10

# Use timestamp in filenames to avoid collisions
_chart_ts = str(int(time.time()))
"""
    full_code = preamble + "\n" + code

    try:
        has_chart = "plt.savefig" in code or "savefig" in code

        with tempfile.NamedTemporaryFile(mode="w", suffix=".py", delete=False) as f:
            f.write(full_code)
            filepath = f.name

        env = os.environ.copy()
        env["MPLBACKEND"] = "Agg"

        try:
            result = subprocess.run(
                ["python3", filepath],
                capture_output=True, text=True, timeout=60, env=env,
            )
        finally:
            try:
                os.unlink(filepath)
            except OSError:
                pass

        output_file = None
        filename = None

        if has_chart:
            # Find charts created during this run (recent files)
            charts = glob.glob("/tmp/*.png")
            recent = [c for c in charts if os.path.getmtime(c) > time.time() - 120]
            if recent:
                recent.sort(key=lambda x: os.path.getmtime(x), reverse=True)
                with open(recent[0], "rb") as f:
                    output_file = base64.b64encode(f.read()).decode()
                filename = os.path.basename(recent[0])
                # Clean up all generated charts
                for c in recent:
                    try:
                        os.unlink(c)
                    except OSError:
                        pass

        return {
            "success": result.returncode == 0,
            "stdout": result.stdout[:10000],
            "stderr": result.stderr[:5000],
            "error": None if result.returncode == 0 else f"Exit code {result.returncode}",
            "output_file": output_file,
            "filename": filename,
        }

    except subprocess.TimeoutExpired:
        return {"success": False, "error": "Analytics timed out (60s)"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ─────────────────────────────────────────────
# Document Parsing and OCR Extraction
# ─────────────────────────────────────────────

@app.function(image=image, timeout=60)
def parse_document(file_b64: str, filename: str):
    import base64
    import tempfile
    import os
    
    file_bytes = base64.b64decode(file_b64)
    ext = os.path.splitext(filename)[1].lower()
    
    try:
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f:
            f.write(file_bytes)
            temp_path = f.name
            
        text_content = ""
        metadata = {}
        
        if ext == ".pdf":
            import pdfplumber
            with pdfplumber.open(temp_path) as pdf:
                pages_text = []
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    # Try to extract structured table data for invoices
                    tables = page.extract_tables()
                    if tables:
                        table_str = "\n--- Structured Table Data on Page " + str(i+1) + " ---\n"
                        for row in tables:
                            table_str += " | ".join([str(cell or "").strip() for cell in row]) + "\n"
                        page_text += table_str
                    pages_text.append(f"--- Page {i+1} ---\n{page_text}")
                text_content = "\n\n".join(pages_text)
                metadata = {"pages": len(pdf.pages)}
        elif ext in [".txt", ".csv", ".json", ".xml", ".yaml", ".yml", ".md"]:
            text_content = file_bytes.decode("utf-8", errors="replace")
        elif ext == ".docx":
            from docx import Document
            doc = Document(temp_path)
            paragraphs = [p.text for p in doc.paragraphs]
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    tables_text.append(" | ".join([cell.text.strip() for cell in row.cells]))
            text_content = "\n".join(paragraphs) + "\n\n--- Tables ---\n" + "\n".join(tables_text)
        else:
            text_content = file_bytes.decode("utf-8", errors="replace")
            
        try:
            os.unlink(temp_path)
        except OSError:
            pass
            
        return {
            "success": True,
            "text": text_content,
            "metadata": metadata
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to parse document: {str(e)}"
        }


@app.function(image=image, timeout=30, secrets=[webhook_secret])
def generate_embedding(text: str) -> dict:
    try:
        from fastembed import TextEmbedding
        model = TextEmbedding()
        embeddings = list(model.embed([text]))
        if len(embeddings) > 0:
            return {
                "success": True,
                "embedding": embeddings[0].tolist()
            }
        else:
            return {
                "success": False,
                "error": "No embeddings returned"
            }
    except Exception as e:
        return {
            "success": False,
            "error": f"Embedding generation failed: {str(e)}"
        }


# ─────────────────────────────────────────────
# Unified Entry Point
# ─────────────────────────────────────────────

@app.function(image=image, timeout=120, secrets=[webhook_secret])
@modal.fastapi_endpoint(method="POST")
async def execute(request: Request):
    # Verify auth
    if not verify_request(request):
        raise HTTPException(status_code=401, detail="Invalid or missing webhook secret")

    payload = await request.json()
    task_type = payload.get("type", "code")

    if task_type == "code":
        return execute_code.remote(
            payload["code"],
            payload.get("language", "python"),
            payload.get("dependencies", "")
        )
    elif task_type == "web_read":
        return read_web_page.remote(payload["url"])
    elif task_type == "document":
        return generate_document.remote(payload)
    elif task_type == "analytics":
        return run_analytics.remote(payload)
    elif task_type == "parse_document":
        return parse_document.remote(payload["file"], payload["filename"])
    elif task_type == "generate_embedding":
        return generate_embedding.remote(payload["text"])
    else:
        return {"success": False, "error": f"Unknown task type: {task_type}"}


def _safe_filename(name: str) -> str:
    import re
    return re.sub(r"[^a-z0-9_-]", "_", name.lower())[:50]
